import os
import docx
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_math_element(xml_str):
    return parse_xml(f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">{xml_str}</m:oMath>')

def mr(text, italic=True):
    sty = '' if italic else '<m:rPr><m:sty m:val="p"/></m:rPr>'
    # Escape XML
    text_clean = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    return f'<m:r>{sty}<m:t>{text_clean}</m:t></m:r>'

def mf(num, den):
    return f'<m:f><m:num>{num}</m:num><m:den>{den}</m:den></m:f>'

def msup(base, sup):
    return f'<m:sSup><m:e>{base}</m:e><m:sup>{sup}</m:sup></m:sSup>'

def msub(base, sub):
    return f'<m:sSub><m:e>{base}</m:e><m:sub>{sub}</m:sub></m:sSub>'

def mdelim(inner, beg='(', end=')'):
    return f'<m:d><m:dPr><m:begChr m:val="{beg}"/><m:endChr m:val="{end}"/></m:dPr><m:e>{inner}</m:e></m:d>'

def meqarr(*rows):
    e_rows = ''.join(f'<m:e>{r}</m:e>' for r in rows)
    return f'<m:eqArr>{e_rows}</m:eqArr>'

def set_cell_border(cell, **kwargs):
    """
    kwargs: top, bottom, left, right
    values: dict(sz=12, val='single', color='000000', space='0')
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            for key, val in edge_data.items():
                element.set(qn(f'w:{key}'), str(val))
            tcBorders.append(element)
        else:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            element.set(qn('w:val'), 'none')
            tcBorders.append(element)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def build_mathplus_exam_docx(output_path):
    doc = docx.Document()
    
    # Page setup - A4, margins 1.8cm
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    
    # Configure Normal Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11.5)
    font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.space_before = Pt(0)
    
    # ----------------------------------------------------
    # HEADER CONFIGURATION
    # ----------------------------------------------------
    header = section.header
    header.is_linked_to_previous = False
    
    htable = header.add_table(rows=1, cols=2, width=Cm(17.0))
    htable.alignment = WD_TABLE_ALIGNMENT.CENTER
    htable.autofit = False
    
    c0 = htable.cell(0, 0)
    c0.width = Cm(8.5)
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_logo = c0.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_logo.paragraph_format.space_after = Pt(0)
    p_logo.paragraph_format.space_before = Pt(0)
    logo_path = '/Volumes/data/code/mathplus/public/images/logo.jpg'
    if os.path.exists(logo_path):
        run_img = p_logo.add_run()
        run_img.add_picture(logo_path, height=Cm(1.1))
    
    c1 = htable.cell(0, 1)
    c1.width = Cm(8.5)
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_title = c1.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_title.paragraph_format.space_after = Pt(0)
    p_title.paragraph_format.space_before = Pt(0)
    run_htext = p_title.add_run("TRUNG TÂM MATHPLUS ACADEMY")
    run_htext.bold = True
    run_htext.font.name = 'Times New Roman'
    run_htext.font.size = Pt(12)
    run_htext.font.color.rgb = RGBColor(0, 0, 0)
    
    for row in htable.rows:
        for cell in row.cells:
            set_cell_border(cell)
            
    # ----------------------------------------------------
    # FOOTER CONFIGURATION
    # ----------------------------------------------------
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_after = Pt(0)
    fp.paragraph_format.space_before = Pt(0)
    run_foot = fp.add_run("Nguyễn Dũng -0977961189.")
    run_foot.bold = True
    run_foot.font.name = 'Times New Roman'
    run_foot.font.size = Pt(10.5)
    run_foot.font.color.rgb = RGBColor(0, 0, 0)
    
    # ----------------------------------------------------
    # TOP HEADER BLOCK (Exam Title Table)
    # ----------------------------------------------------
    top_table = doc.add_table(rows=1, cols=2)
    top_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    top_table.autofit = False
    top_table.columns[0].width = Cm(7.5)
    top_table.columns[1].width = Cm(9.5)
    
    # Left Cell
    cell_left = top_table.cell(0, 0)
    cell_left.width = Cm(7.5)
    set_cell_border(cell_left)
    p_l1 = cell_left.paragraphs[0]
    p_l1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_l1.paragraph_format.space_after = Pt(0)
    r_l1 = p_l1.add_run("PHÒNG GIÁO DỤC VÀ ĐÀO TẠO\n")
    r_l1.bold = True
    r_l1.font.size = Pt(11)
    r_l2 = p_l1.add_run("HUYỆN GIA LÂM")
    r_l2.bold = True
    r_l2.font.size = Pt(11)
    
    # Underline short border or space
    p_box_wrap = cell_left.add_paragraph()
    p_box_wrap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_box_wrap.paragraph_format.space_before = Pt(6)
    p_box_wrap.paragraph_format.space_after = Pt(0)
    
    # Sub table for boxed "Đề thi gồm 2 trang"
    box_tbl = cell_left.add_table(rows=1, cols=1)
    box_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    box_tbl.autofit = False
    box_cell = box_tbl.cell(0, 0)
    box_cell.width = Cm(4.2)
    set_cell_margins(box_cell, top=40, bottom=40, left=80, right=80)
    border_box = dict(sz=6, val='single', color='000000', space='0')
    set_cell_border(box_cell, top=border_box, bottom=border_box, left=border_box, right=border_box)
    p_b_in = box_cell.paragraphs[0]
    p_b_in.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_b_in.paragraph_format.space_after = Pt(0)
    p_b_in.paragraph_format.space_before = Pt(0)
    r_box = p_b_in.add_run("Đề thi gồm 2 trang")
    r_box.bold = True
    r_box.font.size = Pt(10)
    
    # Right Cell
    cell_right = top_table.cell(0, 1)
    cell_right.width = Cm(9.5)
    set_cell_border(cell_right)
    p_r1 = cell_right.paragraphs[0]
    p_r1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r1.paragraph_format.space_after = Pt(0)
    r_r1 = p_r1.add_run("ĐỀ KHẢO SÁT CHẤT LƯỢNG THÁNG 9\n")
    r_r1.bold = True
    r_r1.font.size = Pt(11.5)
    r_r2 = p_r1.add_run("NĂM HỌC 2024-2025\n")
    r_r2.bold = True
    r_r2.font.size = Pt(11.5)
    r_r3 = p_r1.add_run("MÔN: TOÁN 9\n")
    r_r3.bold = True
    r_r3.font.size = Pt(11.5)
    r_r4 = p_r1.add_run("Thời gian làm bài: 120 phút")
    r_r4.italic = True
    r_r4.font.size = Pt(11)
    
    # Space after header block
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_after = Pt(6)
    p_sep.paragraph_format.space_before = Pt(0)
    
    # ----------------------------------------------------
    # BÀI I
    # ----------------------------------------------------
    p_b1 = doc.add_paragraph()
    p_b1.paragraph_format.space_after = Pt(2)
    r_b1 = p_b1.add_run("Bài I. (1,5 điểm)")
    r_b1.bold = True
    r_b1.italic = True
    
    # 1)
    p_b1_1 = doc.add_paragraph()
    p_b1_1.paragraph_format.space_after = Pt(3)
    p_b1_1.paragraph_format.line_spacing = 1.15
    r1 = p_b1_1.add_run("1) ")
    r1.bold = True
    p_b1_1.add_run("Khối 9 của một trường THCS tổ chức giải bóng đá với bốn đội tham dự là các đội bóng của các lớp A, B, C và D. Trước giải đấu, câu lạc bộ Thể dục thể thao đã thực hiện một cuộc khảo sát kín dự đoán của các thành viên về đội bóng sẽ vô địch giải đấu và thu được kết quả như sau:")
    
    # Table 2x10
    tbl1 = doc.add_table(rows=2, cols=10)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl1.autofit = False
    
    row1_data = ["A", "B", "A", "A", "A", "A", "A", "B", "D", "B"]
    row2_data = ["A", "A", "B", "D", "D", "A", "A", "B", "D", "D"]
    
    for c_idx in range(10):
        tbl1.columns[c_idx].width = Cm(1.3)
        
    border_spec = dict(sz=4, val='single', color='000000', space='0')
    
    for r_idx, r_data in enumerate([row1_data, row2_data]):
        for c_idx, val in enumerate(r_data):
            cell = tbl1.cell(r_idx, c_idx)
            cell.width = Cm(1.3)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=40, bottom=40, left=40, right=40)
            set_cell_border(cell, top=border_spec, bottom=border_spec, left=border_spec, right=border_spec)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(val)
            r.bold = True
            r.font.size = Pt(11)
            
    p_b1_1_sub = doc.add_paragraph()
    p_b1_1_sub.paragraph_format.space_before = Pt(3)
    p_b1_1_sub.paragraph_format.space_after = Pt(4)
    p_b1_1_sub.add_run("Lập bảng thống kê về số lượng dự đoán vô địch cho mỗi đội.")
    
    # 2) Lucky wheel side-by-side
    p_b1_2_title = doc.add_paragraph()
    p_b1_2_title.paragraph_format.space_after = Pt(2)
    p_b1_2_title.paragraph_format.space_before = Pt(2)
    r2_num = p_b1_2_title.add_run("2) ")
    r2_num.bold = True
    r2_title = p_b1_2_title.add_run("Trò chơi vòng quay may mắn.")
    r2_title.italic = True
    
    # Side-by-side table for wheel problem
    tbl_wheel = doc.add_table(rows=1, cols=2)
    tbl_wheel.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_wheel.autofit = False
    tbl_wheel.columns[0].width = Cm(12.0)
    tbl_wheel.columns[1].width = Cm(5.0)
    
    c_w_left = tbl_wheel.cell(0, 0)
    c_w_left.width = Cm(12.0)
    c_w_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(c_w_left)
    
    p_wl1 = c_w_left.paragraphs[0]
    p_wl1.paragraph_format.space_before = Pt(0)
    p_wl1.paragraph_format.space_after = Pt(3)
    p_wl1.paragraph_format.line_spacing = 1.15
    p_wl1.add_run("Một bánh xe hình tròn được chia thành 12 hình quạt như nhau, trong đó có 2 hình quạt ghi 100 điểm, 2 hình quạt ghi 200 điểm, 2 hình quạt ghi 300 điểm, 2 hình quạt ghi 400 điểm, và 1 hình quạt ghi 500 điểm, 2 hình quạt ghi 1000 điểm và 1 hình quạt ghi 2000 điểm như hình bên.")
    
    p_wl2 = c_w_left.add_paragraph()
    p_wl2.paragraph_format.space_before = Pt(3)
    p_wl2.paragraph_format.space_after = Pt(0)
    p_wl2.paragraph_format.line_spacing = 1.15
    p_wl2.add_run("Bạn Lan chơi trò chơi này. Tính xác suất của biến cố A: “Trong một lượt quay, Lan được 1000 điểm.”")
    
    c_w_right = tbl_wheel.cell(0, 1)
    c_w_right.width = Cm(5.0)
    c_w_right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(c_w_right)
    p_wr = c_w_right.paragraphs[0]
    p_wr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_wr.paragraph_format.space_before = Pt(0)
    p_wr.paragraph_format.space_after = Pt(0)
    wheel_img_path = '/Volumes/data/code/mathplus/scratch/pdf_extract/lucky_wheel_perfect.png'
    if os.path.exists(wheel_img_path):
        p_wr.add_run().add_picture(wheel_img_path, width=Cm(4.0))
        
    p_sep2 = doc.add_paragraph()
    p_sep2.paragraph_format.space_after = Pt(2)
    p_sep2.paragraph_format.space_before = Pt(0)
    
    # ----------------------------------------------------
    # BÀI II
    # ----------------------------------------------------
    p_b2 = doc.add_paragraph()
    p_b2.paragraph_format.space_after = Pt(3)
    r_b2 = p_b2.add_run("Bài II. (1,5 điểm) ")
    r_b2.bold = True
    r_b2.italic = True
    p_b2.add_run("Cho hai biểu thức: ")
    
    # Formula A = (x^2 - 4)/x
    f_A = mr('A') + mr(' = ') + mf(msup(mr('x'), mr('2')) + mr(' - 4'), mr('x'))
    p_b2._p.append(create_math_element(f_A))
    
    p_b2.add_run(" và ")
    
    # Formula B = 3/(x-2) + (2x+3)/(4-x^2)
    f_B = mr('B') + mr(' = ') + mf(mr('3'), mr('x - 2')) + mr(' + ') + mf(mr('2x + 3'), mr('4 - ') + msup(mr('x'), mr('2')))
    p_b2._p.append(create_math_element(f_B))
    
    p_b2.add_run(" với ")
    f_cond = mr('x ≠ 0; x ≠ ±2', italic=False)
    p_b2._p.append(create_math_element(f_cond))
    p_b2.add_run(".")
    
    # 1)
    p_b2_1 = doc.add_paragraph()
    p_b2_1.paragraph_format.space_after = Pt(2)
    r = p_b2_1.add_run("1) ")
    r.bold = True
    p_b2_1.add_run("Tính giá trị của biểu thức ")
    p_b2_1._p.append(create_math_element(mr('A')))
    p_b2_1.add_run(" khi ")
    p_b2_1._p.append(create_math_element(mr('x = 4')))
    p_b2_1.add_run(".")
    
    # 2)
    p_b2_2 = doc.add_paragraph()
    p_b2_2.paragraph_format.space_after = Pt(2)
    r = p_b2_2.add_run("2) ")
    r.bold = True
    p_b2_2.add_run("Chứng minh rằng ")
    f_B_res = mr('B') + mr(' = ') + mf(mr('x + 3'), msup(mr('x'), mr('2')) + mr(' - 4'))
    p_b2_2._p.append(create_math_element(f_B_res))
    p_b2_2.add_run(".")
    
    # 3)
    p_b2_3 = doc.add_paragraph()
    p_b2_3.paragraph_format.space_after = Pt(4)
    r = p_b2_3.add_run("3) ")
    r.bold = True
    p_b2_3.add_run("Xét biểu thức ")
    p_b2_3._p.append(create_math_element(mr('P') + mr(' = ') + mr('A') + mr(' ⋅ ') + mr('B')))
    p_b2_3.add_run(". Tìm ")
    p_b2_3._p.append(create_math_element(mr('x')))
    p_b2_3.add_run(" để ")
    f_P_val = mr('P') + mr(' = ') + mf(mr('2'), mr('3'))
    p_b2_3._p.append(create_math_element(f_P_val))
    p_b2_3.add_run(".")
    
    # ----------------------------------------------------
    # BÀI III (Part 1 & 2 on Page 1)
    # ----------------------------------------------------
    p_b3 = doc.add_paragraph()
    p_b3.paragraph_format.space_after = Pt(2)
    r_b3 = p_b3.add_run("Bài III. (2,5 điểm)")
    r_b3.bold = True
    r_b3.italic = True
    
    # 1)
    p_b3_1 = doc.add_paragraph()
    p_b3_1.paragraph_format.space_after = Pt(3)
    p_b3_1.paragraph_format.line_spacing = 1.15
    r = p_b3_1.add_run("1) ")
    r.bold = True
    p_b3_1.add_run("Cô Linh chia số tiền 500 triệu đồng của mình cho hai khoản đầu tư. Sau một năm, tổng số tiền lãi thu được là 28 triệu đồng. Lãi suất cho khoản đầu tư thứ nhất là 5%/năm và khoản đầu tư thứ hai là 6%/năm. Tính số tiền cô Linh đầu tư cho mỗi khoản.")
    
    # 2)
    p_b3_2 = doc.add_paragraph()
    p_b3_2.paragraph_format.space_after = Pt(3)
    p_b3_2.paragraph_format.line_spacing = 1.15
    r = p_b3_2.add_run("2) ")
    r.bold = True
    p_b3_2.add_run("Một tổ sản xuất theo kế hoạch mỗi ngày phải sản xuất 40 sản phẩm. Khi thực hiện, nhờ cải tiến kỹ thuật nên mỗi ngày tổ sản xuất được 45 sản phẩm. Do đó không những tổ hoàn thành trước kế hoạch 1 ngày mà còn vượt mức 10 sản phẩm. Hỏi theo kế hoạch tổ phải sản xuất bao nhiêu sản phẩm?")
    
    # ----------------------------------------------------
    # PAGE BREAK (Page 2 begins with Bài III. 3)
    # ----------------------------------------------------
    doc.add_page_break()
    
    # 3) on Page 2
    p_b3_3 = doc.add_paragraph()
    p_b3_3.paragraph_format.space_before = Pt(4)
    p_b3_3.paragraph_format.space_after = Pt(4)
    r = p_b3_3.add_run("3) ")
    r.bold = True
    p_b3_3.add_run("Tìm các số ")
    p_b3_3._p.append(create_math_element(mr('a')))
    p_b3_3.add_run(" và ")
    p_b3_3._p.append(create_math_element(mr('b')))
    p_b3_3.add_run(" biết hệ phương trình ")
    
    sys_eq = mdelim(meqarr(mr('ax + by = 3'), mr('2ax - 3by = 1')), beg='{', end='')
    p_b3_3._p.append(create_math_element(sys_eq))
    p_b3_3.add_run(" có nghiệm ")
    sol_pt = mdelim(mr('x; y'), beg='(', end=')') + mr(' = ') + mdelim(mr('2; 1'), beg='(', end=')')
    p_b3_3._p.append(create_math_element(sol_pt))
    p_b3_3.add_run(".")
    
    # ----------------------------------------------------
    # BÀI IV
    # ----------------------------------------------------
    p_b4 = doc.add_paragraph()
    p_b4.paragraph_format.space_after = Pt(2)
    r_b4 = p_b4.add_run("Bài IV. (4,0 điểm)")
    r_b4.bold = True
    r_b4.italic = True
    
    # 1)
    p_b4_1 = doc.add_paragraph()
    p_b4_1.paragraph_format.space_after = Pt(1)
    r = p_b4_1.add_run("1)")
    r.bold = True
    
    # Slope diagram centered
    slope_img_path = '/Volumes/data/code/mathplus/scratch/pdf_extract/slope_perfect.png'
    if os.path.exists(slope_img_path):
        p_slope_img = doc.add_paragraph()
        p_slope_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_slope_img.paragraph_format.space_before = Pt(2)
        p_slope_img.paragraph_format.space_after = Pt(4)
        p_slope_img.add_run().add_picture(slope_img_path, width=Cm(8.0))
        
    p_b4_1_txt = doc.add_paragraph()
    p_b4_1_txt.paragraph_format.space_after = Pt(3)
    p_b4_1_txt.paragraph_format.line_spacing = 1.15
    p_b4_1_txt.add_run("Ta có thể xác định góc dốc ")
    p_b4_1_txt._p.append(create_math_element(mr('α')))
    p_b4_1_txt.add_run(" của một đoạn đường dốc khi biết độ dài của dốc là ")
    p_b4_1_txt._p.append(create_math_element(mr('a')))
    p_b4_1_txt.add_run(" và độ cao của đỉnh dốc so với đường nằm ngang là ")
    p_b4_1_txt._p.append(create_math_element(mr('h')))
    p_b4_1_txt.add_run(". Trong các tòa chung cư, người ta thường thiết kế đoạn dốc cho người đi xe lăn với góc dốc nhỏ hơn ")
    p_b4_1_txt._p.append(create_math_element(mr('6°', italic=False)))
    p_b4_1_txt.add_run(". Ở một tòa chung cư, cho biết đoạn dốc vào sảnh tòa nhà dài 4 m và độ cao của đỉnh dốc là 0,4 m.")
    
    p_b4_1_a = doc.add_paragraph()
    p_b4_1_a.paragraph_format.left_indent = Cm(0.6)
    p_b4_1_a.paragraph_format.space_after = Pt(2)
    r = p_b4_1_a.add_run("a) ")
    r.bold = True
    p_b4_1_a.add_run("Hãy tính góc dốc.")
    
    p_b4_1_b = doc.add_paragraph()
    p_b4_1_b.paragraph_format.left_indent = Cm(0.6)
    p_b4_1_b.paragraph_format.space_after = Pt(4)
    r = p_b4_1_b.add_run("b) ")
    r.bold = True
    p_b4_1_b.add_run("Hỏi góc đó có đúng tiêu chuẩn của dốc cho người đi xe lăn không?")
    
    # 2)
    p_b4_2 = doc.add_paragraph()
    p_b4_2.paragraph_format.space_after = Pt(2)
    r = p_b4_2.add_run("2) ")
    r.bold = True
    p_b4_2.add_run("Cho tam giác ")
    p_b4_2._p.append(create_math_element(mr('ABC')))
    p_b4_2.add_run(" vuông tại ")
    p_b4_2._p.append(create_math_element(mr('A')))
    p_b4_2.add_run(", vẽ đường cao ")
    p_b4_2._p.append(create_math_element(mr('AH')))
    p_b4_2.add_run(".")
    
    # a)
    p_b4_2_a = doc.add_paragraph()
    p_b4_2_a.paragraph_format.left_indent = Cm(0.6)
    p_b4_2_a.paragraph_format.space_after = Pt(2)
    r = p_b4_2_a.add_run("a) ")
    r.bold = True
    p_b4_2_a.add_run("Tính độ dài các cạnh ")
    p_b4_2_a._p.append(create_math_element(mr('AC') + mr(', ') + mr('AH')))
    p_b4_2_a.add_run(" và số đo góc ")
    p_b4_2_a._p.append(create_math_element(mr('B')))
    p_b4_2_a.add_run(" (làm tròn đến phút) nếu ")
    p_b4_2_a._p.append(create_math_element(mr('AB = 3') + mr(' cm') + mr('; ') + mr('BC = 5') + mr(' cm')))
    p_b4_2_a.add_run(".")
    
    # b)
    p_b4_2_b = doc.add_paragraph()
    p_b4_2_b.paragraph_format.left_indent = Cm(0.6)
    p_b4_2_b.paragraph_format.space_after = Pt(2)
    p_b4_2_b.paragraph_format.line_spacing = 1.15
    r = p_b4_2_b.add_run("b) ")
    r.bold = True
    p_b4_2_b.add_run("Qua ")
    p_b4_2_b._p.append(create_math_element(mr('H')))
    p_b4_2_b.add_run(" kẻ các đường thẳng vuông góc với ")
    p_b4_2_b._p.append(create_math_element(mr('AB')))
    p_b4_2_b.add_run(" và ")
    p_b4_2_b._p.append(create_math_element(mr('AC')))
    p_b4_2_b.add_run(" lần lượt tại ")
    p_b4_2_b._p.append(create_math_element(mr('D')))
    p_b4_2_b.add_run(" và ")
    p_b4_2_b._p.append(create_math_element(mr('E')))
    p_b4_2_b.add_run(". Chứng minh tứ giác ")
    p_b4_2_b._p.append(create_math_element(mr('ADHE')))
    p_b4_2_b.add_run(" là hình chữ nhật và ")
    p_b4_2_b._p.append(create_math_element(mr('AD ⋅ AB = AE ⋅ AC')))
    p_b4_2_b.add_run(" từ đó suy ra ")
    p_b4_2_b._p.append(create_math_element(mr('ΔABC ∼ ΔAED')))
    p_b4_2_b.add_run(" (hoặc ")
    p_b4_2_b._p.append(create_math_element(mr('ΔABC')))
    p_b4_2_b.add_run(" đồng dạng với ")
    p_b4_2_b._p.append(create_math_element(mr('ΔAED')))
    p_b4_2_b.add_run(").")
    
    # c)
    p_b4_2_c = doc.add_paragraph()
    p_b4_2_c.paragraph_format.left_indent = Cm(0.6)
    p_b4_2_c.paragraph_format.space_after = Pt(4)
    p_b4_2_c.paragraph_format.line_spacing = 1.15
    r = p_b4_2_c.add_run("c) ")
    r.bold = True
    p_b4_2_c.add_run("Kẻ ")
    p_b4_2_c._p.append(create_math_element(mr('AI ⊥ DE')))
    p_b4_2_c.add_run(" (")
    p_b4_2_c._p.append(create_math_element(mr('I ∈ DE')))
    p_b4_2_c.add_run("), ")
    p_b4_2_c._p.append(create_math_element(mr('AI')))
    p_b4_2_c.add_run(" cắt ")
    p_b4_2_c._p.append(create_math_element(mr('BC')))
    p_b4_2_c.add_run(" tại ")
    p_b4_2_c._p.append(create_math_element(mr('M')))
    p_b4_2_c.add_run(". Chứng minh ")
    p_b4_2_c._p.append(create_math_element(mr('M')))
    p_b4_2_c.add_run(" là trung điểm của ")
    p_b4_2_c._p.append(create_math_element(mr('BC')))
    p_b4_2_c.add_run(".")
    
    # ----------------------------------------------------
    # BÀI V
    # ----------------------------------------------------
    p_b5 = doc.add_paragraph()
    p_b5.paragraph_format.space_after = Pt(2)
    r_b5 = p_b5.add_run("Bài V. (0,5 điểm)")
    r_b5.bold = True
    r_b5.italic = True
    
    p_b5_txt = doc.add_paragraph()
    p_b5_txt.paragraph_format.space_after = Pt(4)
    p_b5_txt.paragraph_format.line_spacing = 1.15
    p_b5_txt.add_run("Người ta trộn 4 kg chất lỏng loại I với 3 kg chất lỏng loại II thì được một hỗn hợp có khối lượng riêng là 700 kg/m³")
    p_b5_txt.add_run(". Biết khối lượng riêng của chất lỏng loại I lớn hơn khối lượng riêng của chất lỏng loại II 200 kg/m³")
    p_b5_txt.add_run(". Tính khối lượng riêng của mỗi chất.")
    
    # ----------------------------------------------------
    # END SECTION & SIGNATURE
    # ----------------------------------------------------
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_end.paragraph_format.space_before = Pt(8)
    p_end.paragraph_format.space_after = Pt(2)
    r_end = p_end.add_run("---------------------- Hết ----------------------")
    r_end.bold = True
    
    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_note.paragraph_format.space_after = Pt(6)
    r_note = p_note.add_run("Cán bộ coi thi không giải thích gì thêm")
    r_note.italic = True
    r_note.font.size = Pt(10.5)
    
    # Signature table for perfect side-by-side alignment
    sig_tbl = doc.add_table(rows=1, cols=2)
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_tbl.autofit = False
    sig_tbl.columns[0].width = Cm(10.5)
    sig_tbl.columns[1].width = Cm(6.5)
    
    c_s0 = sig_tbl.cell(0, 0)
    c_s0.width = Cm(10.5)
    set_cell_border(c_s0)
    p_s0 = c_s0.paragraphs[0]
    p_s0.paragraph_format.space_before = Pt(4)
    p_s0.paragraph_format.space_after = Pt(0)
    r_s0_1 = p_s0.add_run("Họ và tên thí sinh: ")
    r_s0_1.bold = True
    r_s0_1.italic = True
    p_s0.add_run("..................................................")
    
    c_s1 = sig_tbl.cell(0, 1)
    c_s1.width = Cm(6.5)
    set_cell_border(c_s1)
    p_s1 = c_s1.paragraphs[0]
    p_s1.paragraph_format.space_before = Pt(4)
    p_s1.paragraph_format.space_after = Pt(0)
    r_s1_1 = p_s1.add_run("Số báo danh: ")
    r_s1_1.bold = True
    r_s1_1.italic = True
    p_s1.add_run("...........................")
    
    # Save
    doc.save(output_path)
    print(f"Generated {output_path} successfully!")

if __name__ == '__main__':
    out_file = '/Volumes/data/code/mathplus/documents/de_mathplus_2_mat.docx'
    build_mathplus_exam_docx(out_file)
